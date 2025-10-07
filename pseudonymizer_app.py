# pseudonymizer_app.py
"""Streamlit app for legal document pseudonymization."""
import streamlit as st
import pandas as pd
from pathlib import Path
import time
from io import BytesIO

import docx
import PyPDF2
import pdfplumber
from docx import Document

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from pseudonymscript import PseudonymizationPipeline, PseudonymConfig, pseudonymize_text

st.set_page_config(
    page_title="Document Pseudonymizer",
    page_icon="🔒",
    layout="wide"
)
 
class DocumentProcessor:
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        try:
            text = file.read().decode('utf-8')
        except UnicodeDecodeError:
            file.seek(0)
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    file.seek(0)
                    text = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file")
        
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        
        return text
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        doc = Document(file)
        full_text = []
        
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text = '\n'.join(full_text)
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        
        return text
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        text_content = []
        
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
        except Exception:
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        text = '\n'.join(text_content)
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        
        return text
    
    @classmethod
    def process_file(cls, uploaded_file) -> tuple[str, str]:
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        if file_extension == '.txt':
            content = cls.extract_text_from_txt(uploaded_file)
        elif file_extension == '.docx':
            content = cls.extract_text_from_docx(uploaded_file)
        elif file_extension == '.pdf':
            content = cls.extract_text_from_pdf(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        content = content.replace('"', '"').replace('"', '"')
        content = content.replace(''', "'").replace(''', "'")
        
        return content, uploaded_file.name


def create_pdf(content: str, filename: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
    )
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    normal_style.leading = 12
    
    story = []
    
    title = Paragraph(f"Pseudonymized Document: {filename}", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            cleaned_para = para.replace('\n', ' ').strip()
            if cleaned_para:
                p = Paragraph(cleaned_para, normal_style)
                story.append(p)
                story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def create_mapping_pdf(mapping: dict) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
    )
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    normal_style.leading = 12
    
    story = []
    
    title = Paragraph("Entity Replacement Mapping", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    for original, replacement in mapping.items():
        entry = Paragraph(f"<b>{original}</b> → {replacement}", normal_style)
        story.append(entry)
        story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def init_session():
    if 'original_content' not in st.session_state:
        st.session_state.original_content = ""
    if 'pseudonymized_content' not in st.session_state:
        st.session_state.pseudonymized_content = ""
    if 'replacement_mapping' not in st.session_state:
        st.session_state.replacement_mapping = {}
    if 'filename' not in st.session_state:
        st.session_state.filename = ""
    if 'processing_time' not in st.session_state:
        st.session_state.processing_time = 0.0


def get_entity_type(replacement):
    if "Plaintiff" in replacement or "Defendant" in replacement or "Attorney" in replacement or "Counsel" in replacement:
        return "Legal Person"
    elif replacement.startswith("Person"):
        return "Person"
    elif replacement.startswith("ORG") or replacement.startswith("Bank"):
        return "Company"
    elif replacement.startswith("Country"):
        return "Country"
    elif replacement.startswith("State"):
        return "State"
    elif replacement.startswith("City"):
        return "City"
    elif replacement.startswith("Building"):
        return "Building"
    elif replacement.startswith("[ADDRESS"):
        return "Address"
    elif "USD" in replacement or "EUR" in replacement or "GBP" in replacement:
        return "Money"
    elif any(month in replacement for month in ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]):
        return "Date"
    else:
        return "Other"


def show_user_guide():
    st.title("User Guide")
    
    st.markdown("## What this tool does")
    st.markdown("""
    The Legal Document Pseudonymizer identifies and replaces sensitive information in legal documents with consistent alternatives.
    
    - Share documents while protecting confidentiality
    - Maintain document structure and readability
    - Keep consistent references (the same person is always "Plaintiff A")
    - Support compliance requirements for data privacy
    """)
    
    st.markdown("---")
    
    st.markdown("## How to Use")
    st.markdown("""
    ### Step 1: Upload Document
    Go to the Document Processor tab and upload your file (TXT, DOCX, or PDF)
    
    ### Step 2: Automatic Processing
    The document processes automatically
    
    ### Step 3: Review Results
    Compare original vs pseudonymized text side-by-side
    
    ### Step 4: Check Entity Mappings
    See what was changed in the mapping table
    
    ### Step 5: Download
    Get your pseudonymized document as PDF or text file
    """)
    
    st.markdown("---")
    
    st.markdown("## What Gets Replaced")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### People with Legal Roles")
        st.markdown("- Plaintiffs, defendants, attorneys")
        st.markdown("- Example: Anna Lee, Director (\"Plaintiff\") → Plaintiff A")
        
        st.markdown("### People (Generic)")
        st.markdown("- Names without legal roles")
        st.markdown("- Example: John Smith → Person A")
        
        st.markdown("### Companies")
        st.markdown("- Business names with corporate suffixes")
        st.markdown("- Example: Orion Holdings Ltd → ORG A Ltd")
        
        st.markdown("### Places")
        st.markdown("- Countries, states, cities")
        st.markdown("- Example: Singapore → Country A")
    
    with col2:
        st.markdown("### Money")
        st.markdown("- Currency amounts (randomized ±15%)")
        st.markdown("- Example: USD 450,000 → USD 466,529")
        
        st.markdown("### Dates")
        st.markdown("- Dates with interval preservation")
        st.markdown("- Example: 3 May 2021 → 21 February 2000")
        
        st.markdown("### Addresses")
        st.markdown("- Street addresses (hash-based codes)")
        st.markdown("- Example: One Raffles Quay → [ADDRESS 963332]")
        
        st.markdown("### Numbers")
        st.markdown("- Percentages, quantities (randomized)")
        st.markdown("- Example: 45% → 52%")
    
    st.markdown("---")
    
    st.markdown("## Sample Entity Mapping")
    
    sample_data = {
        "Type": ["Legal Person", "Legal Person", "Legal Person", "Person", "Company", "Company", "Location", "Money", "Date", "Address"],
        "Original": [
            "Anna Lee, Director of Orion Holdings Ltd (\"Plaintiff\")",
            "Carlos Rivera (\"Defendant\")",
            "Attorney Jason Tan",
            "Anna Lee",
            "Orion Holdings Ltd",
            "DBS Bank Ltd, Raffles Place Branch",
            "Singapore",
            "USD 450,000",
            "3 May 2021",
            "One Raffles Quay"
        ],
        "Replaced With": [
            "Plaintiff A",
            "Defendant B",
            "Attorney A",
            "Plaintiff A (cross-reference)",
            "ORG A Ltd",
            "Bank B Ltd, Location B Branch",
            "Country T",
            "USD 466,529",
            "21 February 2000",
            "[ADDRESS 963332]"
        ]
    }
    
    sample_df = pd.DataFrame(sample_data)
    st.dataframe(sample_df, use_container_width=True, hide_index=True)
    
    st.markdown("---")
    
    st.markdown("## Important Notes")
    
    st.info("""
    Consistency: The same entity will always get the same replacement throughout the document.
    
    Legal Role Priority: People with legal roles (Plaintiff, Defendant, Attorney) are prioritized over generic person names.
    
    Cross-References: If "Anna Lee" is introduced as "Plaintiff", later mentions of just "Anna Lee" will also be replaced with "Plaintiff A".
    
    Date Intervals: When dates are pseudonymized, the time intervals between them are preserved.
    
    Money Randomization: Amounts are randomized by ±15% to prevent reverse-engineering while maintaining realistic values.
    
    Review Required: Always review the output to ensure sensitive information is properly handled.
    
    Compliance: This tool helps with data privacy but does not guarantee complete compliance. Legal review is recommended.
    """)


def show_document_processor():
    st.title("Document Processor")
    
    uploaded_file = st.file_uploader(
        "Choose your document",
        type=['txt', 'docx', 'pdf'],
        help="Upload TXT, DOCX, or PDF files"
    )
    
    if uploaded_file is not None:
        try:
            content, filename = DocumentProcessor.process_file(uploaded_file)
            st.session_state.original_content = content
            st.session_state.filename = filename
            
            st.success(f"File loaded: {filename} ({len(content):,} characters)")

            with st.spinner("Processing document..."):
                config = PseudonymConfig()
                
                start_time = time.time()
                pseudonymized_content, replacement_mapping = pseudonymize_text(content, config)
                processing_time = time.time() - start_time
                
                st.session_state.pseudonymized_content = pseudonymized_content
                st.session_state.replacement_mapping = replacement_mapping
                st.session_state.processing_time = processing_time
            
            st.success("Document processed successfully!")
                
        except Exception as e:
            st.error(f"Error: {str(e)}")
    
    if st.session_state.replacement_mapping or st.session_state.pseudonymized_content:
        
        st.markdown("---")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Entities Found", len(st.session_state.replacement_mapping))
        with col2:
            st.metric("Replacements Made", len(st.session_state.replacement_mapping))
        with col3:
            st.metric("Processing Time", f"{st.session_state.processing_time:.1f}s")
        
        st.markdown("### Document Comparison")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Original**")
            original_text = st.text_area(
                "original",
                value=st.session_state.original_content,
                height=400,
                disabled=False,
                label_visibility="collapsed"
            )
        
        with col2:
            st.markdown("**Pseudonymized**")
            pseudo_text = st.text_area(
                "pseudonymized", 
                value=st.session_state.pseudonymized_content,
                height=400,
                disabled=False,
                label_visibility="collapsed"
            )
        
        if st.session_state.replacement_mapping:
            st.markdown("---")
            st.markdown("### Entity Mapping")
            
            mapping_data = []
            for original, replacement in st.session_state.replacement_mapping.items():
                entity_type = get_entity_type(replacement)
                
                mapping_data.append({
                    "Type": entity_type,
                    "Original": original,
                    "Replaced With": replacement
                })
            
            df = pd.DataFrame(mapping_data)
            st.dataframe(
                df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Type": st.column_config.TextColumn("Type", width="small"),
                    "Original": st.column_config.TextColumn("Original Text", width="large"),
                    "Replaced With": st.column_config.TextColumn("Replacement", width="medium")
                }
            )
        
        st.markdown("---")
        st.markdown("### Download")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            pdf_buffer = create_pdf(st.session_state.pseudonymized_content, st.session_state.filename)
            st.download_button(
                label="Download as PDF",
                data=pdf_buffer.getvalue(),
                file_name=f"pseudonymized_{Path(st.session_state.filename).stem}.pdf",
                mime="application/pdf"
            )
        
        with col2:
            st.download_button(
                label="Download as Text",
                data=st.session_state.pseudonymized_content,
                file_name=f"pseudonymized_{st.session_state.filename}",
                mime="text/plain"
            )
        
        with col3:
            if st.session_state.replacement_mapping:
                mapping_pdf = create_mapping_pdf(st.session_state.replacement_mapping)
                st.download_button(
                    label="Download Mapping PDF",
                    data=mapping_pdf.getvalue(),
                    file_name="entity_mapping.pdf",
                    mime="application/pdf"
                )
            else:
                st.button("No Mapping Available", disabled=True)


def show_entity_mapping_table():
    st.title("Entity Mapping Table")
    
    if not st.session_state.replacement_mapping:
        st.info("Please process a document first on the Document Processor tab to see the entity mappings.")
        return
    
    st.markdown(f"**Document:** {st.session_state.filename}")
    st.markdown(f"**Total Entities Replaced:** {len(st.session_state.replacement_mapping)}")
    
    st.markdown("---")
    
    if st.session_state.replacement_mapping:
        mapping_data = []
        for original, replacement in st.session_state.replacement_mapping.items():
            entity_type = get_entity_type(replacement)
            
            mapping_data.append({
                "Type": entity_type,
                "Original": original,
                "Replaced With": replacement,
                "Length": len(original)
            })
        
        df = pd.DataFrame(mapping_data)
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            entity_types = ["All"] + sorted(df["Type"].unique().tolist())
            selected_type = st.selectbox("Filter by Type:", entity_types)
            
            if selected_type != "All":
                df = df[df["Type"] == selected_type]
        
        with col2:
            search_term = st.text_input("Search original text:", placeholder="Enter text to search...")
            if search_term:
                df = df[df["Original"].str.contains(search_term, case=False, na=False)]
        
        st.markdown(f"**Showing {len(df)} of {len(mapping_data)} entities**")
        
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Type": st.column_config.TextColumn("Type", width="small"),
                "Original": st.column_config.TextColumn("Original Text", width="medium"),
                "Replaced With": st.column_config.TextColumn("Replacement", width="medium"),
                "Length": st.column_config.NumberColumn("Characters", width="small")
            }
        )
        
        if len(df) > 0:
            st.markdown("---")
            st.markdown("### Summary Statistics")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total Entities", len(df))
            
            with col2:
                avg_length = df["Length"].mean()
                st.metric("Avg. Length", f"{avg_length:.1f}")
            
            with col3:
                max_length = df["Length"].max()
                longest = df[df["Length"] == max_length]["Original"].iloc[0]
                st.metric("Longest Text", f"{max_length} chars")
                st.caption(f'"{longest[:30]}..."' if len(longest) > 30 else f'"{longest}"')
            
            with col4:
                type_counts = df["Type"].value_counts()
                most_common = type_counts.index[0]
                st.metric("Most Common Type", type_counts.iloc[0])
                st.caption(most_common)
        
        st.markdown("---")
        st.markdown("### Export Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            csv_data = df.to_csv(index=False)
            st.download_button(
                label="Download as CSV",
                data=csv_data,
                file_name="entity_mapping.csv",
                mime="text/csv"
            )
        
        with col2:
            json_data = df.to_json(orient="records", indent=2)
            st.download_button(
                label="Download as JSON",
                data=json_data,
                file_name="entity_mapping.json",
                mime="application/json"
            )
        
        with col3:
            if len(df) > 0:
                mapping_dict = dict(zip(df["Original"], df["Replaced With"]))
                mapping_pdf = create_mapping_pdf(mapping_dict)
                st.download_button(
                    label="Download as PDF",
                    data=mapping_pdf.getvalue(),
                    file_name="entity_mapping_filtered.pdf",
                    mime="application/pdf"
                )


def main():
    init_session()
    
    st.title("Legal Document Pseudonymizer")
    
    tab1, tab2, tab3 = st.tabs(["User Guide", "Document Processor", "Entity Mapping Table"])
    
    with tab1:
        show_user_guide()
    
    with tab2:
        show_document_processor()
    
    with tab3:
        show_entity_mapping_table()


if __name__ == "__main__":
    main()